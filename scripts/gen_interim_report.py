#!/usr/bin/env python3
"""Generate the project Interim Report as a .docx file."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ACCENT = RGBColor(0x1F, 0x4E, 0x79)   # deep blue
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for lvl, sz in ((1, 15), (2, 12.5)):
    h = doc.styles[f"Heading {lvl}"]
    h.font.name = "Calibri"
    h.font.size = Pt(sz)
    h.font.color.rgb = ACCENT
    h.font.bold = True


def add_bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


# ============================ Title block ============================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Interim Report")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("An AI Tutoring System for the Hong Kong Diploma of Secondary Education (HKDSE)")
r.font.size = Pt(13)
r.font.color.rgb = GREY

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Built on the open-source DeepTutor platform")
r.italic = True
r.font.size = Pt(10.5)
r.font.color.rgb = GREY

# small info line (placeholders for the team to fill)
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = info.add_run("Project Group: ____________     Members: ____________     Supervisor: ____________     Date: ____________")
r.font.size = Pt(9.5)
r.font.color.rgb = GREY

doc.add_paragraph()

# ============================ 1. Introduction ============================
doc.add_heading("1. Introduction", level=1)
doc.add_paragraph(
    "This project builds a focused AI tutoring system for the Hong Kong Diploma of Secondary "
    "Education (HKDSE). Rather than starting from scratch, we build on DeepTutor, an open-source, "
    "agent-based tutoring platform from the HKU Data Intelligence Lab. The platform already "
    "provides a chat workspace, autonomous tutor agents, a knowledge base (RAG) system, and "
    "configurable model providers. Our contribution is a new \"Market\" module that turns the "
    "platform's general AI ability into concrete, exam-oriented tools for three HKDSE subjects: "
    "Chinese, English, and Mathematics."
)
doc.add_paragraph(
    "The goal of this interim stage was to get a complete end-to-end pipeline working — from "
    "configuring the underlying models, to importing real past papers, to producing useful study "
    "outputs such as generated papers, graded essays, and step-by-step solution checking."
)

# ============================ 2. Objectives ============================
doc.add_heading("2. Project Objectives", level=1)
add_bullet(" Integrate a dedicated HKDSE module into the DeepTutor platform.", "Integration:")
add_bullet(" Provide subject-specific tools for Chinese, English, and Mathematics.", "Coverage:")
add_bullet(" Ground question generation in real HKDSE past papers using retrieval (RAG).", "Grounding:")
add_bullet(" Grade student work against the official HKDSE marking rubrics.", "Assessment:")
add_bullet(" Keep the system flexible across cloud and (planned) local model providers.", "Flexibility:")

# ============================ 3. Work completed / preliminary results ============================
doc.add_heading("3. Work Completed and Preliminary Results", level=1)
doc.add_paragraph(
    "The core pipeline is now fully working and has been tested end-to-end. The main achievements "
    "in this period are summarised below."
)

doc.add_heading("3.1 Platform integration and model configuration", level=2)
add_bullet(" Merged the custom HKDSE module onto the latest DeepTutor release (v1.4.2) and resolved "
           "the integration and compatibility issues that arose from the merge.")
add_bullet(" Configured the language model (DeepSeek) and the embedding model (Jina, 1024-dim) "
           "through the platform's model catalog, and verified live connectivity.")

doc.add_heading("3.2 Knowledge bases from real past papers", level=2)
add_bullet(" Imported 2012 HKDSE past papers (Chinese, English, Mathematics) into three separate "
           "knowledge bases, embedded via Jina.")
add_bullet(" All three knowledge bases initialise successfully and are available for retrieval-"
           "augmented generation.")

doc.add_heading("3.3 HKDSE feature set (working end-to-end)", level=2)
doc.add_paragraph(
    "Three subjects are live, each with dedicated tools. The following were verified to run "
    "successfully end-to-end:"
)
add_bullet(" Paper generation grounded in the past-paper knowledge base; essay grading on the "
           "official three-dimension rubric; classical-Chinese sentence analysis.", "Chinese — ")
add_bullet(" Paper generation (including the DSE Summary Writing question type); essay coaching on "
           "the Content / Language / Organisation rubric; an Integrated Skills simulator.", "English — ")
add_bullet(" Line-by-line solution step checking with error localisation; syllabus-based topic "
           "drills; past-paper-grounded paper generation.", "Mathematics — ")

doc.add_heading("3.4 Robustness fix", level=2)
add_bullet(" Identified and fixed a retrieval bug in which the BM25 ranking step failed on small "
           "knowledge bases; the retriever now safely adapts to the available corpus size.")

# ============================ 4. Current limitations ============================
doc.add_heading("4. Current Limitations", level=1)
add_bullet(" Several past papers are scanned image PDFs with little extractable text, so the "
           "retrieved context is currently thin. OCR is needed to fully use this material.", "Document quality: ")
add_bullet(" Retrieval quality and chunking are still basic and need tuning for exam content.", "RAG quality: ")
add_bullet(" The system currently depends on cloud model APIs; local inference is not yet supported.", "Deployment: ")
add_bullet(" Each subject offers a first set of tools; the product design can be deepened with more "
           "meaningful, pedagogically-driven learning features.", "Product depth: ")

# ============================ 5. Timeline ============================
doc.add_heading("5. Timeline for the Remaining Work", level=1)
doc.add_paragraph(
    "The remaining work is organised into the milestones below so that progress can be monitored "
    "closely. (Months are indicative and can be aligned to the course schedule.)"
)

rows = [
    ("Phase", "Planned Work", "Target"),
    ("M1 — RAG optimisation",
     "Add OCR for scanned papers; improve chunking and retrieval quality; tune prompts so generated "
     "content is better grounded in real past papers.",
     "Month 1"),
    ("M2 — Local inference",
     "Support local model providers (e.g. Ollama / vLLM / LM Studio) for both LLM and embedding, so "
     "the system can run without cloud APIs.",
     "Month 2"),
    ("M3 — Subject product design",
     "Enrich each subject's toolset and design more meaningful learning features (e.g. progress "
     "tracking, targeted practice, feedback loops).",
     "Months 2-3"),
    ("M4 — Evaluation",
     "Evaluate output quality against marking rubrics and past-paper standards; collect feedback and "
     "iterate.",
     "Month 3"),
    ("M5 — Final report & demo",
     "Polish the system, finalise documentation, and prepare the final report and demonstration.",
     "Month 4"),
]

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
for i, text in enumerate(rows[0]):
    hdr[i].text = ""
    p = hdr[i].paragraphs[0]
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10.5)
for phase, work, target in rows[1:]:
    cells = table.add_row().cells
    cells[0].text = phase
    cells[1].text = work
    cells[2].text = target
    for c in cells:
        for p in c.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)

# set rough column widths
for row in table.rows:
    row.cells[0].width = Inches(1.7)
    row.cells[1].width = Inches(3.8)
    row.cells[2].width = Inches(1.0)

doc.add_paragraph()

# ============================ 6. Conclusion ============================
doc.add_heading("6. Conclusion", level=1)
doc.add_paragraph(
    "At the interim stage, the full pipeline is working: the platform is integrated, models are "
    "configured, real past papers are imported, and all three HKDSE subjects can generate papers, "
    "grade work, and check solutions end-to-end. The next phase focuses on optimising retrieval, "
    "adding local inference, and deepening the subject-level product design with more meaningful "
    "learning features, following the timeline above."
)

out = "/Users/henry.han/HKU Project/DeepTutor/HKDSE_Interim_Report.docx"
doc.save(out)
print("saved:", out)
